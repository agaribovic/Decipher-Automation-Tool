from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph


def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within *parent*,
    in document order. Each returned value is an instance of either Table or
    Paragraph. *parent* would most commonly be a reference to a main
    Document object, but also works for a _Cell object, which itself can
    contain paragraphs and tables.
    """
    if isinstance(parent, _Document):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def table_print(block):

    table = block
    row_number = 0

    for row in table.rows:

        row_number += 1
        cell_number = 0
        for cell in row.cells:
            join_para = ""
            cell_number += 1
            for paragraph in cell.paragraphs:
                join_para += " "+paragraph.text
            join_para += "\t"
            print("cell" + str(row_number) + str(cell_number) + "=" + join_para + "\n")
        # docOutput.add_paragraph(joinRow)
        # print(joinPara + "\n")
        # print("\n")
        # docOutput.add_paragraph("\n")
        # y.write("\n")


def delete_paragraph(paragraph):

    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def number_candidate(candidate, helper):

    candidate = candidate.replace(".", "")
    candidate = candidate.replace("\t", "")
    candidate = candidate.replace(" ", "")

    if candidate.lower() in helper.alphabet:
        candidate = str(helper.alphabet[candidate.lower()])

    return candidate
